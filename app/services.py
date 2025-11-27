import re
from io import BytesIO
from pathlib import Path
from typing import Dict, Set

from docx import Document


PLACEHOLDER_PATTERN = re.compile(r"\{([^{}]+)\}")


def find_placeholders(docx_path: Path) -> Set[str]:
    """Parse docx and return set of placeholders {KEY}."""
    doc = Document(docx_path)
    texts = []

    for paragraph in doc.paragraphs:
        texts.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    texts.append(paragraph.text)

    combined = "\n".join(texts)
    matches = PLACEHOLDER_PATTERN.findall(combined)
    return set(m.strip() for m in matches if m.strip())


def _replace_text(text: str, values: Dict[str, str]) -> str:
    output = text
    for key, value in values.items():
        placeholder = "{" + key + "}"
        output = output.replace(placeholder, value or "")
    return output


def _replace_in_paragraph(paragraph, values: Dict[str, str]) -> None:
    if not paragraph.runs:
        return
    combined = "".join(run.text for run in paragraph.runs)
    replaced = _replace_text(combined, values)
    if replaced != combined:
        paragraph.runs[0].text = replaced
        for run in paragraph.runs[1:]:
            run.text = ""


def render_docx_with_values(template_path: Path, values: Dict[str, str]) -> BytesIO:
    """
    Replace placeholders in docx and return file-like bytes.

    Keeps working even when Word splits placeholders across multiple runs.
    """
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, values)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, values)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
